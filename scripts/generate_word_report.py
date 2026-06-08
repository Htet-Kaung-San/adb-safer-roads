"""
Convert findings_summary.md to a Word document (.docx)
for ADB submission (required deliverable format).

Usage:
    python scripts/generate_word_report.py

Output:
    report/findings_summary.docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
MD_PATH   = ROOT / "report" / "findings_summary.md"
DOCX_PATH = ROOT / "report" / "findings_summary.docx"


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Default body font ─────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    def add_para(text: str, style_name: str = "Normal", bold=False, italic=False,
                 color=None, size=None, align=None, space_before=0, space_after=6):
        p = doc.add_paragraph(style=style_name)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if align:
            p.alignment = align
        # Handle inline **bold** and *italic*
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            else:
                run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = RGBColor(*bytes.fromhex(color))
            if size:
                run.font.size = Pt(size)
        return p

    in_table  = False
    table_rows = []
    in_code   = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ───────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                # Render as indented mono paragraph
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(cl)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                doc.add_paragraph()
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Table rows ────────────────────────────────────────────────────
        if line.startswith("|"):
            table_rows.append(line)
            i += 1
            continue
        else:
            if table_rows:
                # Flush table
                rows = [r for r in table_rows if not re.match(r"^\|\s*[-:]+", r)]
                if rows:
                    ncols = rows[0].count("|") - 1
                    t = doc.add_table(rows=len(rows), cols=ncols)
                    t.style = "Table Grid"
                    for ri, row_text in enumerate(rows):
                        cells = [c.strip() for c in row_text.strip("|").split("|")]
                        for ci, cell_text in enumerate(cells[:ncols]):
                            cell = t.cell(ri, ci)
                            # Strip markdown bold from cell text
                            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                            clean = re.sub(r"\*([^*]+)\*", r"\1", clean)
                            cell.text = clean
                            if ri == 0:
                                set_cell_bg(cell, "1E40AF")
                                run = cell.paragraphs[0].runs
                                if run:
                                    run[0].bold = True
                                    run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                doc.add_paragraph()
                table_rows = []

        # ── Headings ──────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            add_para(text, "Heading 1", bold=True, size=16, space_before=12, space_after=6)
        elif line.startswith("## "):
            text = line[3:].strip()
            add_para(text, "Heading 2", bold=True, size=13, space_before=10, space_after=4)
        elif line.startswith("### "):
            text = line[4:].strip()
            add_para(text, "Heading 3", bold=True, size=11, space_before=8, space_after=3)

        # ── Horizontal rules ──────────────────────────────────────────────
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── Bullet points ─────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    p.add_run(part[1:-1]).italic = True
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)

        # ── Blockquotes ───────────────────────────────────────────────────
        elif line.startswith("> "):
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", line[2:])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # ── Blank lines ───────────────────────────────────────────────────
        elif line.strip() == "":
            pass  # skip blank lines

        # ── Italic/metadata lines (starting with *...*) ───────────────────
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            add_para(line[1:-1], italic=True, color="666666", size=9, space_after=3)

        # ── Normal paragraph ──────────────────────────────────────────────
        else:
            add_para(line, space_after=6)

        i += 1

    # Flush any remaining table
    if table_rows:
        rows = [r for r in table_rows if not re.match(r"^\|\s*[-:]+", r)]
        if rows:
            ncols = rows[0].count("|") - 1
            t = doc.add_table(rows=len(rows), cols=ncols)
            t.style = "Table Grid"
            for ri, row_text in enumerate(rows):
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                for ci, cell_text in enumerate(cells[:ncols]):
                    cell = t.cell(ri, ci)
                    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                    cell.text = clean
                    if ri == 0:
                        set_cell_bg(cell, "1E40AF")

    doc.save(docx_path)
    print(f"Word document saved → {docx_path}")
    print(f"File size: {docx_path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    print(f"Converting {MD_PATH.name} → {DOCX_PATH.name}...")
    md_to_docx(MD_PATH, DOCX_PATH)
    print("Done. Submit report/findings_summary.docx to ADB.")
